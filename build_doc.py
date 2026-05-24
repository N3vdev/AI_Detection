from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size  = Pt(22)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size  = Pt(15)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.size  = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_para(text='', space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    add_inline(p, text)
    return p

def add_numbered(text, num):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    add_inline(p, text)
    return p

def add_inline(para, text):
    """Parse **bold** and `code` inline markers and add runs."""
    import re
    # split on **...** and `...`
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for tok in tokens:
        if tok.startswith('**') and tok.endswith('**'):
            r = para.add_run(tok[2:-2])
            set_font(r, bold=True)
        elif tok.startswith('`') and tok.endswith('`'):
            r = para.add_run(tok[1:-1])
            set_font(r, name='Courier New', size=9.5, color=(0x26, 0x62, 0x30))
        else:
            r = para.add_run(tok)
            set_font(r)
    return para

def add_code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    # light grey shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    r = p.add_run('\n'.join(lines))
    set_font(r, name='Courier New', size=9)
    return p

def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, '2E74B5')
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        set_font(r, bold=True, color=(0xFF, 0xFF, 0xFF))
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    # data rows
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        fill = 'DEEAF1' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            shade_cell(cell, fill)
            cell.paragraphs[0].clear()
            add_inline(cell.paragraphs[0], val)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    # column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return tbl

def add_note_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'FFF3CD')
    pPr.append(shd)
    r = p.add_run('Note:  ')
    set_font(r, bold=True, color=(0x85, 0x64, 0x04))
    r2 = p.add_run(text)
    set_font(r2, color=(0x85, 0x64, 0x04))
    return p

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BFBFBF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# Title
add_heading('AI Product Inspection System', level=1)
p = add_para()
r = p.add_run('Technical Overview & Installation Guide')
set_font(r, size=12, italic=True, color=(0x70, 0x70, 0x70))
hr()

# ── What It Does ──────────────────────────────────────────────────────────────
add_heading('What It Does', level=2)
p = add_para()
add_inline(p, 'An automated real-time inspection system for conveyor belt production lines. Products pass under cameras and the system instantly extracts:')

for item in [
    '**Barcode value** (EAN-13, CODE-128, QR code, and others)',
    '**Brand name** — e.g. Parachute, Amul, Coca-Cola',
    '**Category** — Food / Drink / Snack / Skincare / Haircare / Medicine / Household',
    '**Expiry date** — any printed format',
    '**Manufacture / Packed date**',
    '**Batch number / Lot number**',
]:
    add_bullet(item)

hr()

# ── How It Works ──────────────────────────────────────────────────────────────
add_heading('How It Will Work on the Conveyor Belt', level=2)

add_heading('1 — Product Detection (Trigger)', level=3)
p = add_para()
add_inline(p, 'A lightweight YOLO model (**YOLO11n**) runs on the CPU and watches the trigger camera continuously at full belt speed. Every frame is scanned. When a product enters frame and is confirmed across three consecutive frames, a trigger fires.')

add_heading('2 — Camera Snapshot', level=3)
p = add_para()
add_inline(p, 'The moment the trigger fires, the system captures a snapshot from all cameras simultaneously.')

add_heading('3 — Barcode Fast Path', level=3)
p = add_para()
add_inline(p, 'The system scans all captured frames in parallel for a barcode using the ZBar library, running simultaneously across all cameras. Typical time: **15–30ms**.')
p = add_para()
add_inline(p, 'If any camera has a readable barcode, the result is written to the database immediately and the full AI pipeline never runs. Since most products on a production line have barcodes, this path handles the majority of products almost instantly.')

add_heading('4 — Full AI Pipeline (no barcode found)', level=3)
p = add_para()
add_inline(p, 'For products where barcode scanning fails — damaged barcodes, unlabelled items, angle obstruction — the system runs a five-stage vision pipeline:')

stages = [
    ('Stage A — Barcode retry with preprocessing',
     'A custom barcode detection model (**YOLOv8s**, trained on barcode images) finds the barcode region and crops it. The crop is retried through six preprocessing variants (upscale, sharpen, adaptive threshold, Otsu threshold) to recover partially damaged or low-contrast barcodes.'),
    ('Stage B — Region detection',
     '**YOLO-World v2 (medium)** scans the image for label stickers, nutrition panels, brand logos, and ingredient blocks. Isolating these regions means later steps work on focused crops rather than the full frame — critical for reading small inkjet-printed dates at the edge of a jar or bottle.'),
    ('Stage C — Text reading (OCR)',
     '**EasyOCR** (CRAFT text detector + recognition network) locates and reads every text region on the label. After OCR runs, the bounding boxes of all detected text are used to compute a tight crop around the text-dense area. This crop is passed to the VLM in the next step — effectively zooming in on the label.'),
    ('Stage D — Vision Language Model',
     '**Qwen2.5-VL-3B-Instruct** (3 billion parameters, 4-bit quantized) receives up to four inputs: region crops from Stage B, the tight text-region crop from Stage C, and the sharpest full image. It extracts brand, product name, category, expiry date, manufacture date, and batch number in a single inference call.'),
    ('Stage E — Regex fallback',
     'A pattern-matching pass over the raw OCR text catches any date or batch fields the VLM missed. It recognises a comprehensive list of keywords (EXP, BBD, BEST BEFORE, MFD, MFG, PKD, DOM, and many more) across all common date formats. Only fills empty fields — never overrides the VLM.'),
]
for title, body in stages:
    p = add_para(space_after=2)
    r = p.add_run(title)
    set_font(r, bold=True, color=(0x1F, 0x49, 0x7D))
    p2 = add_para(space_after=8)
    p2.paragraph_format.left_indent = Inches(0.25)
    add_inline(p2, body)

hr()

# ── Inspection Queue ──────────────────────────────────────────────────────────
add_heading('Inspection Queue', level=2)
p = add_para()
add_inline(p, 'Products with barcodes are handled by the fast path and never queue — they complete in under 30ms. Products without barcodes are queued for the full pipeline (the VLM takes 5–15 seconds per product on GPU). The queue holds up to 10 slots; if the belt outruns the AI during a burst, the excess product is skipped with a log warning rather than crashing or consuming unbounded memory.')

hr()

# ── Result Format ─────────────────────────────────────────────────────────────
add_heading('Result Format (when no barcode is found)', level=2)
add_code_block([
    '{',
    '  "barcode": null,',
    '  "brand": "Parachute",',
    '  "product_name": "Men Advanced Aftershower Hair Cream",',
    '  "product_category": "Skincare",',
    '  "expiry_date": null,',
    '  "manufacture_date": "02/26",',
    '  "batch_number": "KK039-R",',
    '  "status": "Complete (Vision)",',
    '  "processing_ms": 8420',
    '}',
])
doc.add_paragraph()
for item in [
    '**Complete (Barcode)** — barcode decoded in fast path',
    '**Complete (Vision)** — no barcode, brand extracted by VLM',
    '**Incomplete** — pipeline ran but could not extract brand or barcode',
]:
    add_bullet(item)

hr()

# ── Demo Web App ──────────────────────────────────────────────────────────────
add_heading('Demo Web App  (for testing accuracy only)', level=2)
p = add_para()
add_inline(p, 'A browser-based UI for testing detection accuracy without the physical conveyor setup. Upload photos of a product from any phone or laptop and get the full inspection result. Runs the exact same AI pipeline as the conveyor system. Useful for validating accuracy on new product types before deploying on the line.')
p = add_para()
add_inline(p, 'Can be shared over the internet using Cloudflare Tunnel — gives a public HTTPS URL with no port forwarding or account required.')

hr()

# ── Hardware ──────────────────────────────────────────────────────────────────
add_heading('Recommended Hardware Requirements  (Tested On)', level=2)
add_table(
    ['Component', 'Minimum', 'Recommended'],
    [
        ['GPU',            '8GB VRAM',    '12GB VRAM'],
        ['RAM',            '16GB',        '32GB'],
        ['CPU',            '6-core',      '8-core+ for multi-camera'],
        ['Storage',        '20GB free',   '50GB+ (snapshots grow)'],
        ['OS',             'Windows 10/11', 'Windows 11'],
        ['CUDA',           '12.x',        '12.6'],
        ['Python Version', '3.13',        '3.13'],
    ],
    col_widths=[1.8, 1.8, 2.4]
)

hr()

# ── Installation ──────────────────────────────────────────────────────────────
add_heading('Installation', level=2)

add_heading('Prerequisites', level=3)
p = add_para()
add_inline(p, 'Before running anything, make sure the following are installed on the machine:')

add_numbered('**Python 3.13** — download from python.org. During install, tick **"Add Python to PATH"**. `pip` is included with Python — no separate install needed.', 1)
add_numbered('**CUDA 12.x drivers** — install the latest NVIDIA driver for the GPU. The CUDA toolkit is not needed separately; PyTorch bundles it.', 2)

add_heading('Step 1 — Extract the Files', level=3)
p = add_para()
add_inline(p, 'Unzip the provided archive into any folder, for example `C:\\AI_Inspector\\`. The folder should contain the Python source files, the batch files, and `cloudflared.exe`.')

add_heading('Step 2 — Install Python Dependencies', level=3)
p = add_para()
add_inline(p, 'This happens automatically. On the first run, `run_demo.bat` detects that dependencies have not been installed yet and runs `pip install -r requirements.txt` before starting anything. Subsequent runs skip this step entirely.')
p = add_para()
add_inline(p, 'If you prefer to install manually beforehand, open a terminal in the extracted folder and run:')
add_code_block(['pip install -r requirements.txt'])
p = add_para()
add_inline(p, 'PyTorch is pulled with CUDA 12.6 support automatically. First-time install takes a few minutes depending on internet speed.')
p = add_para(space_after=3)
add_inline(p, 'On first run, the following model weights are downloaded automatically and cached locally:')
for item in [
    '**Qwen2.5-VL-3B-Instruct** (~7 GB) — vision language model',
    '**EasyOCR CRAFT + recognition models** (~200 MB) — text detection and reading',
    '**YOLO-World v2 medium** (~200 MB) — region detection',
    '**YOLO11n** (~6 MB) — conveyor trigger model',
    '**YOLOv8s barcode detector** — custom model included in the zip, no download needed',
]:
    add_bullet(item)
p = add_para()
add_inline(p, 'After the first run everything is cached and no internet connection is needed.')

add_heading('Step 3 — Run the Demo (Test Mode)', level=3)
p = add_para()
add_inline(p, 'The zip includes a batch file called **`run_demo.bat`**. Double-click it or run it from a terminal. It does two things:')
add_numbered('Starts the AI server on your machine at `http://localhost:8000`', 1)
add_numbered('Starts a Cloudflare Tunnel using the included `cloudflared.exe`, which prints a public URL like `https://abc123.trycloudflare.com`', 2)
p = add_para()
add_inline(p, 'Share that URL with anyone on any device. They open it in a browser, upload photos of a product, and get the inspection result. The AI runs on your machine — they only see the web UI. No source code is shared.')
p = add_para()
add_inline(p, 'Both the server and the tunnel stay running until you close the terminal window.')
add_note_box('The public URL changes every time you restart the tunnel. To keep a fixed URL, a free Cloudflare account with a named tunnel can be configured separately.')

add_heading('Conveyor Belt Live Feed', level=3)
p = add_para()
add_inline(p, 'The live conveyor mode — real-time camera feed, automatic trigger, and database logging — will be released later once the physical camera setup and belt configuration have been finalised and fully tested. Camera sources, resolution, and trigger thresholds will be set in the configuration file included with that release.')

hr()

# ── Python Dependencies ───────────────────────────────────────────────────────
add_heading('Python Dependencies', level=2)
p = add_para()
add_inline(p, 'All dependencies are listed in `requirements.txt` included in the zip. Run `pip install -r requirements.txt` to install everything in one step.')
doc.add_paragraph()
add_table(
    ['Library', 'Purpose'],
    [
        ['`torch`, `torchvision`',       'Deep learning framework (CUDA 12.6 build)'],
        ['`transformers`',               'Loads and runs Qwen2.5-VL vision language model'],
        ['`accelerate`, `bitsandbytes`', '4-bit quantization to reduce VRAM usage'],
        ['`ultralytics`',               'Runs YOLO models (barcode detector, YOLO-World, YOLO11n)'],
        ['`easyocr`',                   'Text detection and reading (CRAFT + recognition network)'],
        ['`opencv-contrib-python`',     'Image processing and fallback barcode decoder'],
        ['`pyzbar`',                    'Primary barcode decoder (ZBar library)'],
        ['`Pillow`, `numpy`',           'Image loading and array operations'],
        ['`fastapi`, `uvicorn`',        'Web server for the demo app'],
        ['`python-multipart`',          'Handles image file uploads in the demo'],
        ['`huggingface-hub`',           'Downloads and caches model weights from HuggingFace'],
        ['`pyyaml`',                    'Configuration file parsing'],
    ],
    col_widths=[2.2, 4.0]
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r'd:\Projects\AI_Detection\AI_Product_Inspection_System.docx'
doc.save(out)
print(f'Saved: {out}')
